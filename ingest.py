"""
Generic ingestion pipeline for RAG.
- Ingests PDF files, plain-text files, and arbitrary web URLs
- Chunks text with configurable size and overlap
- Embeds with sentence-transformers, stores in ChromaDB
- Persists source metadata for re-indexing

Usage:
    python ingest.py                    # Re-ingest all cached sources
    python ingest.py --force-refresh    # Re-fetch URL sources too
"""

import json
import time
import argparse
from pathlib import Path
from typing import Dict, List, Tuple
from datetime import datetime
import logging

import requests
from bs4 import BeautifulSoup
import pdfplumber

from sentence_transformers import SentenceTransformer

from config import (
    RAW_DIR, CHROMA_DIR, CHROMA_COLLECTION,
    EMBED_MODEL, EMBED_DEVICE,
    CHUNK_SIZE, CHUNK_OVERLAP,
    LOG_INGESTION
)

logging.basicConfig(
    level=logging.INFO if LOG_INGESTION else logging.WARNING,
    format="[%(levelname)s] %(message)s"
)
logger = logging.getLogger(__name__)

CUSTOM_SOURCES_FILE = RAW_DIR / "custom_sources.json"


# ============================================================================
# TOKEN COUNTING
# ============================================================================

def count_tokens(text: str) -> int:
    """Approximate token count: ~4 characters per token."""
    return len(text) // 4


# ============================================================================
# CHUNKER: Generic text chunker
# ============================================================================

def chunk_text(text: str, title: str, doc_type: str, url: str = "", extra_meta: Dict = None) -> List[Dict]:
    """
    Chunk text into CHUNK_SIZE token segments with CHUNK_OVERLAP overlap.

    Splits on double newlines (paragraphs), buffers to CHUNK_SIZE, and
    prepends the tail of the previous buffer for overlap.

    Args:
        text: Full document text
        title: Human-readable document name
        doc_type: "pdf", "txt", or "url"
        url: Source URL (if applicable)
        extra_meta: Additional metadata fields to include

    Returns:
        List of dicts: {text, metadata}
    """
    if not text or not text.strip():
        return []

    if extra_meta is None:
        extra_meta = {}

    chunks = []
    paragraphs = text.split("\n\n")
    buffer = ""
    chunk_index = 0

    for para in paragraphs:
        buffer_tokens = count_tokens(buffer)
        para_tokens = count_tokens(para)

        if buffer_tokens + para_tokens <= CHUNK_SIZE:
            buffer += para + "\n\n"
        else:
            if buffer.strip():
                chunks.append({
                    "text": buffer.strip(),
                    "metadata": {
                        "doc_type": doc_type,
                        "source": title,
                        "url": url,
                        "chunk_index": chunk_index,
                        "added_by": "user",
                        **extra_meta,
                    }
                })
                chunk_index += 1
                # Overlap: keep last CHUNK_OVERLAP tokens of previous buffer
                tail_chars = CHUNK_OVERLAP * 4
                overlap_text = buffer[-tail_chars:] if len(buffer) > tail_chars else buffer
                buffer = overlap_text + para + "\n\n"
            else:
                buffer = para + "\n\n"

    # Flush remaining buffer
    if buffer.strip():
        chunks.append({
            "text": buffer.strip(),
            "metadata": {
                "doc_type": doc_type,
                "source": title,
                "url": url,
                "chunk_index": chunk_index,
                "added_by": "user",
                **extra_meta,
            }
        })

    return chunks


# ============================================================================
# INGESTION: PDF
# ============================================================================

def ingest_pdf(file_path: str, title: str, url_hint: str = "") -> Tuple[List[Dict], int]:
    """
    Extract text from a PDF file and chunk it.

    Injects [Page N] markers so page numbers survive into chunks.

    Args:
        file_path: Absolute path to the PDF file
        title: Human-readable document name
        url_hint: Optional URL to associate with the source

    Returns:
        (chunks, page_count)
    """
    try:
        with pdfplumber.open(file_path) as pdf:
            pages_text = []
            for page_num, page in enumerate(pdf.pages, start=1):
                text = page.extract_text()
                if text and text.strip():
                    pages_text.append((page_num, text))

        if not pages_text:
            logger.warning(f"[PDF] No text extracted from {file_path}")
            return [], 0

        full_text = "\n\n".join(f"[Page {p}]\n{t}" for p, t in pages_text)
        chunks = chunk_text(full_text, title, "pdf", url_hint)

        logger.info(f"[PDF] {title}: {len(pages_text)} pages → {len(chunks)} chunks")
        return chunks, len(pages_text)
    except Exception as e:
        logger.error(f"[PDF] Error extracting text from {file_path}: {e}")
        raise ValueError(f"Failed to read PDF: {e}")


# ============================================================================
# INGESTION: Plain text
# ============================================================================

def ingest_txt(file_path: str, title: str, url_hint: str = "") -> Tuple[List[Dict], int]:
    """
    Read a plain-text file and chunk it.

    Args:
        file_path: Absolute path to the text file
        title: Human-readable document name
        url_hint: Optional URL to associate with the source

    Returns:
        (chunks, line_count)
    """
    text = Path(file_path).read_text(encoding="utf-8", errors="replace")
    text = text.replace("\r\n", "\n").replace("\r", "\n")

    line_count = len(text.splitlines())
    chunks = chunk_text(text, title, "txt", url_hint)

    logger.info(f"[TXT] {title}: {line_count} lines → {len(chunks)} chunks")
    return chunks, line_count


# ============================================================================
# INGESTION: Word (.docx)
# ============================================================================

def ingest_docx(file_path: str, title: str, url_hint: str = "") -> Tuple[List[Dict], int]:
    """Extract text from a .docx Word document and chunk it."""
    from docx import Document as DocxDocument

    doc = DocxDocument(file_path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    text = "\n".join(parts)
    chunks = chunk_text(text, title, "docx", url_hint)
    logger.info(f"[DOCX] {title}: {len(parts)} blocks → {len(chunks)} chunks")
    return chunks, len(parts)


# ============================================================================
# INGESTION: Word legacy (.doc)
# ============================================================================

def ingest_doc(file_path: str, title: str, url_hint: str = "") -> Tuple[List[Dict], int]:
    """
    Extract text from a legacy .doc Word document via antiword.
    Requires the `antiword` binary to be installed in the environment.
    """
    import subprocess
    try:
        result = subprocess.run(
            ["antiword", file_path],
            capture_output=True, text=True, timeout=60, check=True
        )
        text = result.stdout
    except FileNotFoundError as e:
        raise RuntimeError("antiword is not installed — cannot process legacy .doc files") from e
    except subprocess.CalledProcessError as e:
        raise RuntimeError(f"antiword failed: {e.stderr.strip() or e}") from e

    line_count = len(text.splitlines())
    chunks = chunk_text(text, title, "doc", url_hint)
    logger.info(f"[DOC] {title}: {line_count} lines → {len(chunks)} chunks")
    return chunks, line_count


# ============================================================================
# INGESTION: Web URL
# ============================================================================

def _extract_text_requests(url: str) -> str:
    """Fallback scraper using requests + BeautifulSoup."""
    response = requests.get(url, timeout=30, headers={"User-Agent": "Mozilla/5.0"})
    response.raise_for_status()
    soup = BeautifulSoup(response.content, "html.parser")
    for tag in soup(["script", "style", "nav", "footer", "header", "aside"]):
        tag.decompose()
    return soup.get_text(separator="\n", strip=True)


def _extract_text_scrapling(url: str) -> str:
    """Primary scraper using Scrapling — handles JS-rendered pages and anti-bot."""
    from scrapling import Fetcher, PlayWrightFetcher
    try:
        # Try fast fetch first (handles most sites + basic anti-bot)
        fetcher = Fetcher(auto_match=False)
        page = fetcher.get(url, timeout=30)
        text = page.get_all_text(ignore_tags=("script", "style", "nav", "footer", "header", "aside"))
        if text and len(text) >= 200:
            logger.info(f"[URL] Scrapling fast fetch succeeded for {url}")
            return text
    except Exception as e:
        logger.warning(f"[URL] Scrapling fast fetch failed ({e}), trying PlayWright...")

    # Fallback to Playwright for JS-heavy sites
    fetcher = PlayWrightFetcher(auto_match=False)
    page = fetcher.get(url, timeout=60)
    return page.get_all_text(ignore_tags=("script", "style", "nav", "footer", "header", "aside"))


def ingest_url(url: str, title: str) -> Tuple[List[Dict], int]:
    """
    Fetch a web page, extract clean text, and chunk it.
    Uses Scrapling (JS rendering + anti-bot) with fallback to requests/BeautifulSoup.

    Args:
        url: Web URL to fetch
        title: Human-readable document name

    Returns:
        (chunks, word_count)

    Raises:
        ValueError: If insufficient text is extracted
    """
    text = ""

    # Try Scrapling first
    try:
        text = _extract_text_scrapling(url)
        logger.info(f"[URL] Scrapling extracted {len(text)} chars from {url}")
    except Exception as e:
        logger.warning(f"[URL] Scrapling failed ({e}), falling back to requests")
        try:
            text = _extract_text_requests(url)
            logger.info(f"[URL] requests fallback extracted {len(text)} chars from {url}")
        except Exception as e2:
            raise ValueError(f"Failed to fetch {url}: {e2}")

    if len(text) < 100:
        raise ValueError(f"Insufficient text extracted from {url} ({len(text)} chars)")

    word_count = len(text.split())
    chunks = chunk_text(text, title, "url", url)

    logger.info(f"[URL] {title}: {word_count} words → {len(chunks)} chunks")
    return chunks, word_count


# ============================================================================
# CHROMADB: Embed and store chunks
# ============================================================================

def embed_and_store(chunks: List[Dict], collection, embedder, doc_id_prefix: str) -> int:
    """
    Embed chunks and upsert into ChromaDB.

    Args:
        chunks: List of {text, metadata} dicts from a chunk_* function
        collection: ChromaDB collection object
        embedder: SentenceTransformer instance
        doc_id_prefix: Prefix for ChromaDB document IDs

    Returns:
        Number of chunks stored
    """
    for chunk in chunks:
        chunk_index = chunk["metadata"]["chunk_index"]
        doc_id = f"{doc_id_prefix}_{chunk_index}"
        embedding = embedder.encode(chunk["text"], convert_to_tensor=False)

        collection.upsert(
            ids=[doc_id],
            documents=[chunk["text"]],
            embeddings=[embedding.tolist()],
            metadatas=[chunk["metadata"]]
        )

    return len(chunks)


# ============================================================================
# SOURCE METADATA PERSISTENCE
# ============================================================================

def load_custom_sources() -> List[Dict]:
    """Load user-added sources from persistent storage."""
    if not CUSTOM_SOURCES_FILE.exists():
        return []

    try:
        content = CUSTOM_SOURCES_FILE.read_text(encoding="utf-8")
        return json.loads(content) if content.strip() else []
    except Exception as e:
        logger.warning(f"[SOURCES] Error loading: {e}")
        return []


def save_custom_source(source: Dict) -> None:
    """
    Add a source to custom_sources.json (idempotent by title).

    Adds 'id' and 'added' fields automatically.
    """
    sources = load_custom_sources()

    source_id = source.get("title", "unknown").lower().replace(" ", "_")
    source["id"] = source_id
    source["added"] = datetime.now().isoformat()

    if any(s["id"] == source_id for s in sources):
        logger.info(f"[SOURCES] Source already exists: {source_id}")
        return

    sources.append(source)
    CUSTOM_SOURCES_FILE.write_text(json.dumps(sources, indent=2), encoding="utf-8")
    logger.info(f"[SOURCES] Saved: {source_id}")


def remove_custom_source(source_id: str) -> None:
    """Remove a user-added source by ID."""
    sources = load_custom_sources()
    updated = [s for s in sources if s["id"] != source_id]

    if len(updated) == len(sources):
        logger.warning(f"[SOURCES] Source not found: {source_id}")
        return

    CUSTOM_SOURCES_FILE.write_text(json.dumps(updated, indent=2), encoding="utf-8")
    logger.info(f"[SOURCES] Removed: {source_id}")


# ============================================================================
# BUILD INDEX: Re-ingest all saved sources
# ============================================================================

def build_index(force_refresh: bool = False):
    """
    Re-ingest all sources recorded in custom_sources.json.

    PDF and TXT sources are re-embedded from their cached files.
    URL sources are re-fetched only if force_refresh=True.

    Args:
        force_refresh: If True, re-fetch URL sources from the web
    """
    logger.info("=" * 70)
    logger.info("RAG - INGESTION PIPELINE")
    logger.info("=" * 70)

    start_time = time.time()
    total_chunks = 0

    import chromadb
    client = chromadb.PersistentClient(path=str(CHROMA_DIR))
    collection = client.get_or_create_collection(name=CHROMA_COLLECTION)

    logger.info(f"Loading embedding model: {EMBED_MODEL}")
    embedder = SentenceTransformer(EMBED_MODEL, device=EMBED_DEVICE)

    sources = load_custom_sources()
    logger.info(f"Sources to re-index: {len(sources)}")

    for source in sources:
        doc_type = source.get("type", "")
        title = source.get("title", "unknown")
        cached_path = source.get("cached_path", "")
        url = source.get("url", "")
        doc_id_prefix = source.get("id", title.lower().replace(" ", "_"))

        try:
            if doc_type == "pdf" and cached_path and Path(cached_path).exists():
                chunks, _ = ingest_pdf(cached_path, title, url)
            elif doc_type == "txt" and cached_path and Path(cached_path).exists():
                chunks, _ = ingest_txt(cached_path, title, url)
            elif doc_type == "url":
                if force_refresh and url:
                    chunks, _ = ingest_url(url, title)
                else:
                    logger.info(f"[SKIP] URL source '{title}' (use --force-refresh to re-fetch)")
                    continue
            else:
                logger.warning(f"[SKIP] Cannot re-index '{title}' (type={doc_type}, cached_path={cached_path})")
                continue

            count = embed_and_store(chunks, collection, embedder, doc_id_prefix)
            total_chunks += count
            logger.info(f"[OK] {title}: {count} chunks stored")

        except Exception as e:
            logger.error(f"[ERROR] Failed to re-index '{title}': {e}")

    elapsed = time.time() - start_time
    logger.info("=" * 70)
    logger.info("INGESTION COMPLETE")
    logger.info("=" * 70)
    logger.info(f"Total chunks stored: {total_chunks}")
    logger.info(f"Time elapsed: {elapsed:.2f}s")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Re-ingest all sources into ChromaDB")
    parser.add_argument("--force-refresh", action="store_true",
                        help="Re-fetch URL sources from the web")
    args = parser.parse_args()

    build_index(force_refresh=args.force_refresh)

# ============================================================================
# WEB CRAWL: BFS across same-domain links (relevance heuristic)
# ============================================================================

from urllib.parse import urljoin, urlparse, urldefrag
from collections import deque

try:
    from protego import Protego
except Exception:
    Protego = None

_SKIP_EXTS = (
    ".pdf", ".jpg", ".jpeg", ".png", ".gif", ".svg", ".webp", ".ico",
    ".zip", ".tar", ".gz", ".7z", ".rar",
    ".mp3", ".mp4", ".webm", ".avi", ".mov", ".wav", ".m4a",
    ".doc", ".docx", ".xls", ".xlsx", ".ppt", ".pptx",
    ".css", ".js", ".woff", ".woff2", ".ttf",
)


def _normalize_url(url: str) -> str:
    url, _ = urldefrag(url)
    p = urlparse(url)
    netloc = p.netloc.lower()
    path = p.path or "/"
    return p._replace(netloc=netloc, path=path).geturl()


def _extract_links(html_text: str, base_url: str):
    soup = BeautifulSoup(html_text, "html.parser")
    out = []
    for a in soup.find_all("a", href=True):
        href = (a.get("href") or "").strip()
        if not href or href.startswith(("javascript:", "mailto:", "tel:", "#")):
            continue
        absolute = urljoin(base_url, href)
        if absolute.startswith(("http://", "https://")):
            out.append(absolute)
    return out


def _is_relevant(link: str, seed: str, same_domain_only: bool) -> bool:
    seed_p = urlparse(seed)
    link_p = urlparse(link)
    if any(link_p.path.lower().endswith(ext) for ext in _SKIP_EXTS):
        return False
    if same_domain_only and link_p.netloc.lower() != seed_p.netloc.lower():
        return False
    return True


def _fetch_page(url: str):
    response = requests.get(url, timeout=20, headers={"User-Agent": "Mozilla/5.0 (RAGCrawler)"})
    response.raise_for_status()
    ct = (response.headers.get("Content-Type") or "").lower()
    if "html" not in ct and "xml" not in ct and "text" not in ct:
        raise ValueError(f"Non-HTML content type: {ct}")
    html_bytes = response.content
    soup = BeautifulSoup(html_bytes, "html.parser")
    for tag in soup(["script", "style", "nav", "footer", "header", "aside"]):
        tag.decompose()
    text = soup.get_text(separator="\n", strip=True)
    return text, html_bytes.decode("utf-8", errors="ignore")


def ingest_crawl(
    seed_url: str,
    title: str,
    max_depth: int = 2,
    max_pages: int = 20,
    same_domain_only: bool = True,
    respect_robots: bool = False,
) -> Tuple[List[Dict], int]:
    """
    BFS-crawl from seed_url. Each visited page contributes chunks tagged
    with its actual page URL. Respects robots.txt; same-domain by default.
    """
    seed_norm = _normalize_url(seed_url)
    rp = None
    if respect_robots and Protego is not None:
        parsed = urlparse(seed_url)
        robots_url = f"{parsed.scheme}://{parsed.netloc}/robots.txt"
        try:
            r = requests.get(robots_url, timeout=10, headers={"User-Agent": "RAGCrawler"})
            if r.status_code == 200:
                rp = Protego.parse(r.text)
        except Exception as e:
            logger.warning(f"[CRAWL] robots.txt fetch failed: {e}")

    def can_fetch(url):
        if rp is None:
            return True
        try:
            return rp.can_fetch(url, "RAGCrawler")
        except Exception:
            return True

    visited = set()
    queue = deque([(seed_norm, 0)])
    pages = []
    total_chars = 0

    while queue and len(pages) < max_pages:
        url, depth = queue.popleft()
        if url in visited:
            continue
        visited.add(url)
        if not can_fetch(url):
            logger.info(f"[CRAWL] robots.txt disallow: {url}")
            continue
        try:
            text, html_text = _fetch_page(url)
        except Exception as e:
            logger.warning(f"[CRAWL] fetch failed {url}: {e}")
            continue
        if len(text) < 100:
            logger.info(f"[CRAWL] thin content, skip: {url}")
            continue
        pages.append((url, text))
        total_chars += len(text)
        logger.info(f"[CRAWL] {url} (depth={depth}, {len(text)} chars, page {len(pages)}/{max_pages})")
        if depth < max_depth:
            for link in _extract_links(html_text, url):
                norm = _normalize_url(link)
                if norm in visited:
                    continue
                if _is_relevant(norm, seed_url, same_domain_only):
                    queue.append((norm, depth + 1))

    if not pages:
        raise ValueError(f"Crawl extracted no content from {seed_url}")

    all_chunks = []
    for page_url, text in pages:
        page_chunks = chunk_text(text, title, "url", page_url, extra_meta={"page_url": page_url})
        all_chunks.extend(page_chunks)

    word_count = total_chars // 5
    logger.info(f"[CRAWL] {title}: {len(pages)} pages → {len(all_chunks)} chunks")
    return all_chunks, word_count
